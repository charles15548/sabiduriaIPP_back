from pypdf import PdfReader
from io import BytesIO
from fastapi import UploadFile
from script.controllers.libro import subirLibro
from docx import Document
import os
import fitz
import re
from fastapi import HTTPException
from script.ml.variables_globales import MIN_TEXTO_POR_PAGINA

async def extraer_texto(archivo: UploadFile) -> str:
    filename = archivo.filename.lower()

    contenido = await archivo.read()

    # --- PDF ---
    if filename.endswith(".pdf"):
        
        reader = PdfReader(BytesIO(contenido))
        texto_total = []
        for page in reader.pages:
            texto = page.extract_text()
            if texto:
                texto_total.append(texto)

        return "\n".join(texto_total)
    elif filename.endswith(".docx"):
        doc = Document(BytesIO(contenido))
        texto_total = []
        for para in doc.paragraphs:
            if para.text.strip():
                texto_total.append(para.text)
        return "\n".join(texto_total)
    else:
        raise ValueError("Formato de archivo no soportado. Usa PDF o DOCX.")

def limpiar_texto_pdf(texto: str) -> str:
    if not texto:
        return ""

    # Eliminar NUL explícitamente (CRÍTICO)
    texto = texto.replace("\x00", "")

    # Normalizar saltos raros
    texto = texto.replace("\r", "\n")

    # Eliminar caracteres de control (excepto \n y \t)
    texto = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", "", texto)

    # Colapsar espacios excesivos
    texto = re.sub(r"\n{3,}", "\n\n", texto)

    return texto.strip()

def contar_texto(texto: str) -> int:
    if not texto:
        return 0
    return len(texto.strip())


def extraer_paginas_pdf(archivo) -> list:
    contenido = archivo.file.read()
    archivo.file.seek(0)
    doc = fitz.open(stream=contenido,filetype="pdf")
    paginas = []

    for page in doc:
        texto = page.get_text(sort=True)
        texto_limpio = limpiar_texto_pdf(texto)

        if not texto_limpio:
            continue
        if contar_texto(texto_limpio) < MIN_TEXTO_POR_PAGINA:
            continue

        paginas.append({
                "pagina": page.number +1,
                "texto": texto_limpio
        })
    print(f"📘 PÁGINAS ÚTILES: {len(paginas)} / {len(doc)}")
    return paginas


def extraer_paginas_word(archivo) -> list:
    try:
        doc = Document(archivo.file)
    except Exception:
        raise HTTPException(status_code=400, detail="Word corrupto")
    
    paginas =[]
    texto_actual = []

    for p in doc.paragraphs:
        if p.text.strip():
            texto_actual.append(p.text.strip())
    texto_final = "\n".join(texto_actual)

    if contar_texto(texto_final) < MIN_TEXTO_POR_PAGINA:
        return []
    paginas.append({
        "pagina":1,
        "texto": texto_final
    })
    archivo.file.seek(0)
    return paginas


def procesarSubida(nombreLibro, contenido):

    extension = os.path.splitext(contenido.filename)[1].lower()


    if extension == ".pdf":
        paginas = extraer_paginas_pdf(contenido)
    elif extension == ".docx":
        paginas = extraer_paginas_word(contenido)
    else:
        return {"msg": "Formato no soportado"}



    subirLibro(nombreLibro, paginas)